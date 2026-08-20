"""
Media & Asset Inspection Engine.

Three extraction paths depending on file type:
  1. Images -> exiftool subprocess (GPS, device, serial numbers, software)
  2. PDFs   -> PyMuPDF metadata (author, creation tool, local file paths)
  3. Office (.docx/.xlsx) -> python-docx/zipfile XML (creator, last-modified-by,
     revision count, template paths)

Like git_engine, this only extracts and yields RawFinding objects —
no scoring or regex matching happens here.
"""

from __future__ import annotations

import json
import subprocess
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Callable, Iterator
from xml.etree import ElementTree as ET

from opsec_scanner.models import RawFinding, SourceType

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".heic", ".tiff", ".webp"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx", ".xlsx", ".pptx"}

# EXIF tags worth surfacing as individual findings — GPS is the highest
# value (physical location deanonymization), everything else helps build
# a device/software fingerprint.
EXIF_TAGS_OF_INTEREST = {
    "GPSLatitude",
    "GPSLongitude",
    "GPSPosition",
    "Model",           # device model
    "Make",            # device manufacturer
    "Software",        # editing software / OS version
    "LensSerialNumber",
    "SerialNumber",
    "OwnerName",
    "Artist",
    "HostComputer",    # sometimes leaks a machine's hostname
}

# Office XML namespaces for docProps parsing
_CORE_NS = {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
            "dc": "http://purl.org/dc/elements/1.1/",
            "dcterms": "http://purl.org/dc/terms/"}
_APP_NS = {"": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}


def _exiftool_available() -> bool:
    try:
        subprocess.run(["exiftool", "-ver"], capture_output=True, check=True, timeout=5)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return False


def scan_image_exif(path: Path) -> Iterator[RawFinding]:
    """
    Runs exiftool -json on a single image and yields a RawFinding for
    every tag in EXIF_TAGS_OF_INTEREST that's present and non-empty.
    """
    try:
        result = subprocess.run(
            ["exiftool", "-json", "-n", str(path)],
            capture_output=True,
            text=True,
            timeout=15,
            check=True,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
        yield RawFinding(
            source_type=SourceType.MEDIA_EXIF,
            raw_text=f"[error running exiftool: {e}]",
            context="exiftool invocation failed",
            origin=str(path),
        )
        return

    try:
        parsed = json.loads(result.stdout)
    except json.JSONDecodeError:
        return

    if not parsed:
        return

    tags = parsed[0]

    # GPS is handled specially: exiftool with -n gives decimal GPSLatitude/
    # GPSLongitude directly, which is exactly what a reverse-geocoder wants.
    lat = tags.get("GPSLatitude")
    lon = tags.get("GPSLongitude")
    if lat is not None and lon is not None:
        yield RawFinding(
            source_type=SourceType.MEDIA_EXIF,
            raw_text=f"{lat},{lon}",
            context="GPS coordinates (lat,lon)",
            origin=str(path),
            metadata={"tag": "GPSPosition", "lat": lat, "lon": lon},
        )

    for tag, value in tags.items():
        if tag in EXIF_TAGS_OF_INTEREST and tag not in ("GPSLatitude", "GPSLongitude") and value:
            yield RawFinding(
                source_type=SourceType.MEDIA_EXIF,
                raw_text=str(value),
                context=f"EXIF tag: {tag}",
                origin=str(path),
                metadata={"tag": tag},
            )


def scan_image_ocr(path: Path) -> Iterator[RawFinding]:
    """
    OCR pass over images to catch secrets rendered as *visible text* —
    terminal screenshots, browser tab titles, whiteboard photos. This is
    distinct from EXIF (invisible metadata) and was in the original design
    but was skipped in the first build — real leaks live in screenshot
    content at least as often as in metadata.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        yield RawFinding(
            source_type=SourceType.MEDIA_EXIF,
            raw_text="[pytesseract/Pillow not installed, skipping OCR scan]",
            context="dependency missing",
            origin=str(path),
        )
        return

    try:
        img = Image.open(path)
        # Without a timeout, a pathological or maliciously crafted image
        # can hang tesseract indefinitely — exiftool already had a 15s
        # guard for the equivalent risk, OCR didn't. Matters more now
        # that media scanning runs in a thread pool: a hung OCR call
        # ties up a worker slot rather than just blocking a single
        # sequential scan.
        text = pytesseract.image_to_string(img, timeout=15)
    except RuntimeError as e:
        # pytesseract raises RuntimeError specifically on timeout
        yield RawFinding(
            source_type=SourceType.MEDIA_EXIF,
            raw_text=f"[OCR timed out or failed: {e}]",
            context="OCR timeout",
            origin=str(path),
        )
        return
    except Exception as e:
        yield RawFinding(
            source_type=SourceType.MEDIA_EXIF,
            raw_text=f"[error running OCR: {e}]",
            context="OCR failed",
            origin=str(path),
        )
        return

    text = text.strip()
    if not text:
        return

    # Yield the full OCR text as one finding so downstream regex/entropy
    # scanning (which operates on raw_text) can find credential-shaped
    # substrings anywhere in it — same treatment as a git patch line.
    yield RawFinding(
        source_type=SourceType.MEDIA_EXIF,
        raw_text=text,
        context="OCR-extracted text from image",
        origin=str(path),
        metadata={"extraction_method": "ocr"},
    )



    """
    Uses PyMuPDF to pull document metadata (author, creator tool, producer,
    and any embedded local file paths in the creator/producer strings).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        yield RawFinding(
            source_type=SourceType.MEDIA_DOC,
            raw_text="[PyMuPDF not installed, skipping PDF scan]",
            context="dependency missing",
            origin=str(path),
        )
        return

    try:
        doc = fitz.open(path)
    except Exception as e:
        yield RawFinding(
            source_type=SourceType.MEDIA_DOC,
            raw_text=f"[error opening PDF: {e}]",
            context="PyMuPDF open failed",
            origin=str(path),
        )
        return

    meta = doc.metadata or {}
    interesting_fields = ("author", "creator", "producer", "subject", "title")

    for field in interesting_fields:
        value = meta.get(field)
        if value:
            yield RawFinding(
                source_type=SourceType.MEDIA_DOC,
                raw_text=str(value),
                context=f"PDF metadata field: {field}",
                origin=str(path),
                metadata={"field": field},
            )

    # Body text — metadata-only scanning misses secrets typed directly
    # into the document (e.g. a leaked password pasted into a shared PDF).
    try:
        body_text = "\n".join(page.get_text() for page in doc)
        body_text = body_text.strip()
        if body_text:
            yield RawFinding(
                source_type=SourceType.MEDIA_DOC,
                raw_text=body_text,
                context="PDF body text",
                origin=str(path),
                metadata={"field": "body_text"},
            )
    except Exception:
        pass

    doc.close()


def scan_office_document(path: Path) -> Iterator[RawFinding]:
    """
    .docx/.xlsx/.pptx are just zip archives of XML. docProps/core.xml has
    creator/lastModifiedBy/revision; docProps/app.xml sometimes has
    TotalTime and the template path used to create the doc.
    """
    try:
        with zipfile.ZipFile(path) as z:
            names = z.namelist()

            if "docProps/core.xml" in names:
                core_xml = z.read("docProps/core.xml")
                root = ET.fromstring(core_xml)

                creator = root.find("dc:creator", _CORE_NS)
                last_modified_by = root.find("cp:lastModifiedBy", _CORE_NS)
                revision = root.find("cp:revision", _CORE_NS)

                if creator is not None and creator.text:
                    yield RawFinding(
                        source_type=SourceType.MEDIA_DOC,
                        raw_text=creator.text,
                        context="Office docProps creator",
                        origin=str(path),
                        metadata={"field": "creator"},
                    )
                if last_modified_by is not None and last_modified_by.text:
                    yield RawFinding(
                        source_type=SourceType.MEDIA_DOC,
                        raw_text=last_modified_by.text,
                        context="Office docProps lastModifiedBy",
                        origin=str(path),
                        metadata={"field": "lastModifiedBy"},
                    )
                if revision is not None and revision.text:
                    yield RawFinding(
                        source_type=SourceType.MEDIA_DOC,
                        raw_text=f"revision count: {revision.text}",
                        context="Office docProps revision (edit history depth)",
                        origin=str(path),
                        metadata={"field": "revision"},
                    )

            if "docProps/app.xml" in names:
                app_xml = z.read("docProps/app.xml")
                root = ET.fromstring(app_xml)
                # app.xml default namespace varies by producer; search without
                # namespace prefix as a fallback for TotalTime and Template.
                for tag in ("TotalTime", "Template", "Application", "AppVersion"):
                    el = root.find(tag) or root.find(f"{{{_APP_NS['']}}}{tag}")
                    if el is not None and el.text:
                        yield RawFinding(
                            source_type=SourceType.MEDIA_DOC,
                            raw_text=el.text,
                            context=f"Office app.xml field: {tag}",
                            origin=str(path),
                            metadata={"field": tag},
                        )

    except (zipfile.BadZipFile, ET.ParseError, KeyError) as e:
        yield RawFinding(
            source_type=SourceType.MEDIA_DOC,
            raw_text=f"[error parsing office document: {e}]",
            context="office doc parse failed",
            origin=str(path),
        )
        return

    # Body text (docx only — python-docx doesn't read xlsx/pptx bodies).
    # Same gap as PDFs: metadata-only scanning misses secrets pasted
    # directly into the document text.
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document

            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            body_text = "\n".join(paragraphs).strip()
            if body_text:
                yield RawFinding(
                    source_type=SourceType.MEDIA_DOC,
                    raw_text=body_text,
                    context="Office document body text",
                    origin=str(path),
                    metadata={"field": "body_text"},
                )
        except Exception:
            pass


def _scan_single_file(path: Path, exiftool_ok: bool) -> list[RawFinding]:
    """Dispatches one file to the right parser. Factored out of
    scan_media_dir so it can run inside a thread pool worker."""
    ext = path.suffix.lower()
    findings: list[RawFinding] = []

    if ext in IMAGE_EXTENSIONS:
        if exiftool_ok:
            findings.extend(scan_image_exif(path))
        else:
            findings.append(
                RawFinding(
                    source_type=SourceType.MEDIA_EXIF,
                    raw_text="[exiftool not found on PATH, skipping image scan]",
                    context="dependency missing",
                    origin=str(path),
                )
            )
        findings.extend(scan_image_ocr(path))
    elif ext in PDF_EXTENSIONS:
        findings.extend(scan_pdf_metadata(path))
    elif ext in DOCX_EXTENSIONS:
        findings.extend(scan_office_document(path))

    return findings


def scan_media_dir(
    dir_path: Path,
    max_workers: int = 4,
    progress_callback: "Callable[[int, int, Path], None] | None" = None,
) -> list[RawFinding]:
    """
    Entrypoint used by main.py. Walks a directory, dispatches each file
    to the appropriate parser based on extension, and returns a flat
    list of RawFinding objects.

    Runs file processing in a thread pool: exiftool/tesseract are
    subprocess calls and PyMuPDF/python-docx parsing releases the GIL
    for most of its work, so this is I/O-bound enough that threads give
    a real speedup without the complexity of a full asyncio rewrite —
    the honest tradeoff given this is a CLI tool, not a long-running
    server where asyncio's benefits would matter more.

    progress_callback(completed_count, total_count, current_path), if
    given, is invoked after each file finishes — main.py uses this to
    drive a rich progress bar instead of the terminal sitting silent
    during a large media directory scan.
    """
    exiftool_ok = _exiftool_available()
    all_paths = [p for p in sorted(dir_path.rglob("*")) if p.is_file()]
    total = len(all_paths)
    findings: list[RawFinding] = []
    completed = 0

    if total == 0:
        return findings

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_path = {executor.submit(_scan_single_file, p, exiftool_ok): p for p in all_paths}
        for future in as_completed(future_to_path):
            path = future_to_path[future]
            completed += 1
            try:
                findings.extend(future.result())
            except Exception as e:
                findings.append(
                    RawFinding(
                        source_type=SourceType.MEDIA_DOC,
                        raw_text=f"[error processing file: {e}]",
                        context="worker thread failed",
                        origin=str(path),
                    )
                )
            if progress_callback:
                progress_callback(completed, total, path)

    return findings
